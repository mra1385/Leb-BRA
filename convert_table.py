import sys
import html # Import the html module for escaping
from docx import Document
from docx.opc.exceptions import PackageNotFoundError

def convert_docx_table_to_html(docx_filename="comparison_Apr2.docx", output_html_filename="table_content.html"):
    """Reads the first table from a .docx file and saves its HTML <table> representation to a file, adding combined data-label attributes for responsive design."""
    try:
        document = Document(docx_filename)
    except PackageNotFoundError:
        print(f"Error: File '{docx_filename}' not found or is not a valid Word document.", file=sys.stderr)
        sys.exit(1)
    except Exception as e:
        print(f"An unexpected error occurred opening the file: {e}", file=sys.stderr)
        sys.exit(1)

    if not document.tables:
        print(f"Error: No tables found in the document '{docx_filename}'.", file=sys.stderr)
        sys.exit(1)

    table = document.tables[0]

    html_table = "<table>\n"
    html_thead = "  <thead>\n"
    html_tbody = "  <tbody>\n"

    headers = []
    is_first_row = True

    for i, row in enumerate(table.rows):
        row_html = "    <tr>\n"
        if is_first_row:
            tag = "th"
            for cell in row.cells:
                cell_text = cell.text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                cell_text = cell_text.replace('\n', '<br>').replace('\r', '')
                headers.append(cell_text) # Store header text
                row_html += f"      <{tag}>{cell_text}</{tag}>\n"
            row_html += "    </tr>\n"
            html_thead += row_html
            html_thead += "  </thead>\n"
        else:
            tag = "td"
            aspect_text_raw = '' # Store the RAW text from the first cell
            for j, cell in enumerate(row.cells):
                # Get raw text first
                raw_cell_text = cell.text.replace('\r', '') # Remove CR
                
                # Store the raw aspect text from the first cell
                if j == 0:
                    aspect_text_raw = raw_cell_text

                # Prepare text for display inside the TD (with escaping and <br>)
                display_cell_text = raw_cell_text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                display_cell_text = display_cell_text.replace('\n', '<br>')

                # Get the corresponding column header (already escaped when stored)
                column_header = headers[j] if j < len(headers) else ''

                # Construct the data-label using RAW aspect text
                if j == 0: # For the first cell, just use the column header
                    label_text = column_header
                else: # For other cells, combine RAW Aspect and Column Header
                    # Use aspect_text_raw which doesn't have &amp;
                    label_text = f"{aspect_text_raw.strip()}\n{column_header}" if aspect_text_raw else column_header # Use newline as separator

                # Escape the final label for the attribute value
                safe_label = html.escape(label_text, quote=True)
                
                # Use the display_cell_text for the content inside the tag
                row_html += f'      <{tag} data-label="{safe_label}">{display_cell_text}</{tag}>\n'
            row_html += "    </tr>\n"
            html_tbody += row_html

        is_first_row = False

    html_tbody += "  </tbody>\n"
    html_table += html_thead
    html_table += html_tbody
    html_table += "</table>"

    try:
        with open(output_html_filename, 'w', encoding='utf-8') as f:
            f.write(html_table)
        print(f"Successfully converted table to {output_html_filename} with combined data-labels.")
    except IOError as e:
        print(f"Error writing to file {output_html_filename}: {e}", file=sys.stderr)
        sys.exit(1)
    except Exception as e:
        print(f"An unexpected error occurred writing the file: {e}", file=sys.stderr)
        sys.exit(1)

if __name__ == "__main__":
    convert_docx_table_to_html() 